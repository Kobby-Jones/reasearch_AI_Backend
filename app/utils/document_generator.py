"""Academic thesis document generator (DOCX / PDF).

Renders a *block* structure into a properly formatted scholarly document:
title page, auto-updating table of contents, numbered headings, statistical
tables, embedded figures with captions, and a references section, with page
numbers in the footer.

The report service builds the block list — interleaving AI-written prose with
deterministically-computed tables and figures — so every number in the document
comes from the analytics engine, never the language model.

Block types
-----------
{"type": "markdown", "text": "...markdown..."}   # ##/### headings, -, 1., **bold**, *italic*
{"type": "heading", "level": 1|2|3, "text": "..."}
{"type": "paragraph", "text": "..."}
{"type": "table", "title": "...", "columns": [...], "rows": [[...]]}
{"type": "figure", "path": "/abs/x.png", "caption": "..."}
{"type": "references", "items": ["...", "..."]}
{"type": "pagebreak"}
"""
from __future__ import annotations

import os
import re
from typing import Iterable


def _ensure_dir(path: str) -> None:
    d = os.path.dirname(path)
    if d:
        os.makedirs(d, exist_ok=True)


_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _parse_markdown_blocks(text: str) -> list[dict]:
    """Turn AI markdown prose into structured blocks the renderers understand."""
    blocks: list[dict] = []
    lines = (text or "").replace("\r\n", "\n").split("\n")
    para: list[str] = []

    def flush_para():
        if para:
            joined = " ".join(s.strip() for s in para if s.strip())
            if joined:
                blocks.append({"type": "paragraph", "text": joined})
            para.clear()

    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            flush_para()
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_para()
            level = min(3, len(m.group(1)) + 1)  # ## -> 2, ### -> 3
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            continue
        mb = re.match(r"^[-*]\s+(.*)$", stripped)
        if mb:
            flush_para()
            blocks.append({"type": "bullet", "text": mb.group(1).strip()})
            continue
        mn = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if mn:
            flush_para()
            blocks.append({"type": "number", "text": mn.group(1).strip()})
            continue
        para.append(stripped)
    flush_para()
    return blocks


def _inline_segments(text: str) -> list[tuple[str, bool, bool]]:
    """Split text into (segment, bold, italic) runs for **bold** / *italic*."""
    segments: list[tuple[str, bool, bool]] = []
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            segments.append((part[2:-2], True, False))
        elif part.startswith("*") and part.endswith("*"):
            segments.append((part[1:-1], False, True))
        else:
            segments.append((part, False, False))
    return segments or [(text, False, False)]


# ----------------------------------------------------------------------------
# DOCX
# ----------------------------------------------------------------------------
def generate_docx(path: str, meta: dict, blocks: list[dict], watermark: bool = False) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _ensure_dir(path)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    def _heading(text, level):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            r.font.name = "Times New Roman"
        return h

    def _add_page_number_footer():
        section = doc.sections[-1]
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(instr); run._r.append(f2)

    # title page
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(meta.get("title", "Research Report"))
    tr.bold = True; tr.font.size = Pt(20); tr.font.name = "Times New Roman"

    if meta.get("subtitle"):
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(meta["subtitle"]); sr.italic = True; sr.font.size = Pt(13)

    for _ in range(2):
        doc.add_paragraph()
    for line in (meta.get("author"), meta.get("institution"), meta.get("field"), meta.get("date")):
        if line:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(line)).font.size = Pt(12)

    if watermark:
        doc.add_paragraph()
        wm = doc.add_paragraph(); wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wr = wm.add_run("Generated with RAI-Core (FREE tier) — upgrade to remove this notice.")
        wr.italic = True; wr.font.size = Pt(9); wr.font.color.rgb = RGBColor(0x94, 0x94, 0x94)

    doc.add_page_break()

    # table of contents (auto field)
    _heading("Table of Contents", 1)
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose \u201CUpdate Field\u201D to build the contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2, placeholder, f3):
        run._r.append(el)
    doc.add_page_break()
    _add_page_number_footer()

    fig_counter = {"n": 0}
    tbl_counter = {"n": 0}

    def render_block(b: dict):
        bt = b.get("type")
        if bt == "pagebreak":
            doc.add_page_break()
        elif bt == "heading":
            _heading(b["text"], min(3, b.get("level", 2)))
        elif bt == "paragraph":
            p = doc.add_paragraph()
            for seg, bold, ital in _inline_segments(b["text"]):
                r = p.add_run(seg); r.bold = bold; r.italic = ital
        elif bt == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            for seg, bold, ital in _inline_segments(b["text"]):
                r = p.add_run(seg); r.bold = bold; r.italic = ital
        elif bt == "number":
            p = doc.add_paragraph(style="List Number")
            for seg, bold, ital in _inline_segments(b["text"]):
                r = p.add_run(seg); r.bold = bold; r.italic = ital
        elif bt == "markdown":
            for sub in _parse_markdown_blocks(b["text"]):
                render_block(sub)
        elif bt == "table":
            tbl_counter["n"] += 1
            cap = doc.add_paragraph()
            cr = cap.add_run(f"Table {tbl_counter['n']}. {b.get('title','')}")
            cr.bold = True; cr.font.size = Pt(11)
            cols = b.get("columns", [])
            rows = b.get("rows", [])
            if cols:
                table = doc.add_table(rows=1, cols=len(cols))
                try:
                    table.style = "Light Grid Accent 1"
                except KeyError:
                    table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, c in enumerate(cols):
                    hdr[i].text = str(c)
                    for par in hdr[i].paragraphs:
                        for rr in par.runs:
                            rr.bold = True; rr.font.size = Pt(10)
                for row in rows:
                    cells = table.add_row().cells
                    for i, val in enumerate(row[: len(cols)]):
                        cells[i].text = "" if val is None else str(val)
                        for par in cells[i].paragraphs:
                            for rr in par.runs:
                                rr.font.size = Pt(10)
            doc.add_paragraph()
        elif bt == "figure":
            p = b.get("path")
            if p and os.path.exists(p):
                fig_counter["n"] += 1
                pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.add_run().add_picture(p, width=Inches(5.6))
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(f"Figure {fig_counter['n']}. {b.get('caption','').replace('Figure: ','')}")
                cr.italic = True; cr.font.size = Pt(10)
                doc.add_paragraph()
        elif bt == "references":
            _heading("References", 1)
            for item in b.get("items", []):
                rp = doc.add_paragraph(item)
                rp.paragraph_format.left_indent = Inches(0.5)
                rp.paragraph_format.first_line_indent = Inches(-0.5)

    for b in blocks:
        render_block(b)

    doc.save(path)
    return path


# ----------------------------------------------------------------------------
# PDF
# ----------------------------------------------------------------------------
def generate_pdf(path: str, meta: dict, blocks: list[dict], watermark: bool = False) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, PageBreak, Image, Table, TableStyle,
    )
    from reportlab.lib.enums import TA_CENTER

    _ensure_dir(path)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Times-Roman",
                          fontSize=11, leading=16, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Times-Bold",
                        fontSize=16, spaceBefore=14, spaceAfter=8, textColor=colors.HexColor("#111827"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Times-Bold",
                        fontSize=13, spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#1f2937"))
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Times-Bold",
                        fontSize=12, spaceBefore=8, spaceAfter=4)
    cap = ParagraphStyle("Cap", parent=body, fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor("#475569"))

    def esc(s) -> str:
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def inline(s: str) -> str:
        s = esc(s)
        s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
        s = re.sub(r"\*(.+?)\*", r"<i>\1</i>", s)
        return s

    flow: list = []
    flow.append(Spacer(1, 2.2 * inch))
    flow.append(Paragraph(esc(meta.get("title", "Research Report")),
                          ParagraphStyle("T", parent=h1, fontSize=22, alignment=TA_CENTER)))
    if meta.get("subtitle"):
        flow.append(Spacer(1, 8))
        flow.append(Paragraph("<i>%s</i>" % esc(meta["subtitle"]),
                              ParagraphStyle("ST", parent=body, fontSize=13, alignment=TA_CENTER)))
    flow.append(Spacer(1, 0.6 * inch))
    for line in (meta.get("author"), meta.get("institution"), meta.get("field"), meta.get("date")):
        if line:
            flow.append(Paragraph(esc(line), ParagraphStyle("M", parent=body, alignment=TA_CENTER)))
    if watermark:
        flow.append(Spacer(1, 0.4 * inch))
        flow.append(Paragraph("<i>Generated with RAI-Core (FREE tier) — upgrade to remove this notice.</i>", cap))
    flow.append(PageBreak())

    fig_n = {"n": 0}
    tbl_n = {"n": 0}

    def render(b: dict):
        bt = b.get("type")
        if bt == "pagebreak":
            flow.append(PageBreak())
        elif bt == "heading":
            flow.append(Paragraph(esc(b["text"]), {1: h1, 2: h2, 3: h3}.get(b.get("level", 2), h2)))
        elif bt == "paragraph":
            flow.append(Paragraph(inline(b["text"]), body))
        elif bt in ("bullet", "number"):
            prefix = "\u2022 " if bt == "bullet" else "\u2014 "
            flow.append(Paragraph(prefix + inline(b["text"]),
                                  ParagraphStyle("LI", parent=body, leftIndent=18)))
        elif bt == "markdown":
            for sub in _parse_markdown_blocks(b["text"]):
                render(sub)
        elif bt == "table":
            tbl_n["n"] += 1
            flow.append(Paragraph("<b>Table %d.</b> %s" % (tbl_n["n"], esc(b.get("title", ""))), body))
            cols = b.get("columns", [])
            rows = b.get("rows", [])
            if cols:
                data = [[Paragraph("<b>%s</b>" % esc(c), ParagraphStyle("th", parent=body, fontSize=9, textColor=colors.white)) for c in cols]]
                for row in rows:
                    data.append([Paragraph(esc("" if v is None else v), ParagraphStyle("td", parent=body, fontSize=9))
                                 for v in row[: len(cols)]])
                tbl = Table(data, repeatRows=1, hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a8a")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                flow.append(tbl)
            flow.append(Spacer(1, 10))
        elif bt == "figure":
            p = b.get("path")
            if p and os.path.exists(p):
                fig_n["n"] += 1
                try:
                    img = Image(p)
                    max_w = 5.6 * inch
                    if img.drawWidth > max_w:
                        ratio = max_w / img.drawWidth
                        img.drawWidth = max_w
                        img.drawHeight *= ratio
                    img.hAlign = "CENTER"
                    flow.append(img)
                    flow.append(Paragraph("Figure %d. %s" % (fig_n["n"], esc(b.get("caption", "").replace("Figure: ", ""))), cap))
                    flow.append(Spacer(1, 10))
                except Exception:
                    pass
        elif bt == "references":
            flow.append(Paragraph("References", h1))
            for item in b.get("items", []):
                flow.append(Paragraph(esc(item),
                                      ParagraphStyle("ref", parent=body, leftIndent=24, firstLineIndent=-24)))

    for b in blocks:
        render(b)

    def _page_number(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Times-Roman", 9)
        canvas.drawCentredString(A4[0] / 2, 0.5 * inch, str(doc_.page))
        canvas.restoreState()

    doc = SimpleDocTemplate(path, pagesize=A4, title=meta.get("title", "Report"),
                            topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch)
    doc.build(flow, onLaterPages=_page_number, onFirstPage=lambda c, d: None)
    return path


# ----------------------------------------------------------------------------
# Backward-compatible thin wrappers (old (heading, body) tuple signature)
# ----------------------------------------------------------------------------
def _tuples_to_blocks(title: str, chapters: Iterable[tuple[str, str]]) -> tuple[dict, list[dict]]:
    blocks: list[dict] = []
    for heading, bodytext in chapters:
        blocks.append({"type": "pagebreak"})
        blocks.append({"type": "heading", "level": 1, "text": heading})
        blocks.append({"type": "markdown", "text": bodytext})
    return {"title": title}, blocks


def generate_docx_simple(path, title, chapters, watermark=False):
    meta, blocks = _tuples_to_blocks(title, chapters)
    return generate_docx(path, meta, blocks, watermark)


def generate_pdf_simple(path, title, chapters, watermark=False):
    meta, blocks = _tuples_to_blocks(title, chapters)
    return generate_pdf(path, meta, blocks, watermark)


# ----------------------------------------------------------------------------
# LaTeX
# ----------------------------------------------------------------------------
_LATEX_SPECIAL = {
    "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
    "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _tex_escape(text: str) -> str:
    if text is None:
        return ""
    out = []
    for ch in str(text):
        out.append(_LATEX_SPECIAL.get(ch, ch))
    return "".join(out)


def _tex_inline(text: str) -> str:
    """Escape, then apply **bold** / *italic* as LaTeX commands."""
    parts = []
    for seg, bold, ital in _inline_segments(text or ""):
        s = _tex_escape(seg)
        if bold:
            s = r"\textbf{%s}" % s
        if ital:
            s = r"\textit{%s}" % s
        parts.append(s)
    return "".join(parts)


_TEX_SECTION = {1: "section", 2: "subsection", 3: "subsubsection"}


def generate_latex(path: str, meta: dict, blocks: list[dict], watermark: bool = False) -> str:
    """Render report blocks to a self-contained LaTeX (.tex) document.

    Produces the same content as the docx/pdf, in a form a researcher can drop
    into Overleaf or compile locally. Figures are referenced by absolute path so
    the file compiles in place; move the figures alongside if relocating.
    """
    lines: list[str] = []
    a = lines.append
    a(r"\documentclass[12pt,a4paper]{article}")
    a(r"\usepackage[utf8]{inputenc}")
    a(r"\usepackage[T1]{fontenc}")
    a(r"\usepackage{graphicx}")
    a(r"\usepackage{booktabs}")
    a(r"\usepackage{longtable}")
    a(r"\usepackage[margin=1in]{geometry}")
    a(r"\usepackage{setspace}")
    a(r"\usepackage{hyperref}")
    if watermark:
        a(r"\usepackage{draftwatermark}")
        a(r"\SetWatermarkText{DRAFT}")
        a(r"\SetWatermarkScale{1.1}")
    a(r"\title{%s}" % _tex_escape(meta.get("title", "Research Report")))
    a(r"\author{%s}" % _tex_escape(meta.get("author", "")))
    a(r"\date{%s}" % _tex_escape(meta.get("date", "")))
    a(r"\onehalfspacing")
    a(r"\begin{document}")
    a(r"\maketitle")
    if meta.get("subtitle"):
        a(r"\begin{center}\textit{%s}\end{center}" % _tex_escape(meta["subtitle"]))
    a(r"\tableofcontents")
    a(r"\newpage")

    fig_n = {"n": 0}
    tbl_n = {"n": 0}

    def render(b: dict) -> None:
        bt = b.get("type")
        if bt == "pagebreak":
            a(r"\newpage")
        elif bt == "heading":
            cmd = _TEX_SECTION.get(min(3, b.get("level", 2)), "section")
            a(r"\%s{%s}" % (cmd, _tex_escape(b.get("text", ""))))
        elif bt == "paragraph":
            a(_tex_inline(b.get("text", "")))
            a("")
        elif bt == "bullet":
            a(r"\begin{itemize}\item %s\end{itemize}" % _tex_inline(b.get("text", "")))
        elif bt == "number":
            a(r"\begin{enumerate}\item %s\end{enumerate}" % _tex_inline(b.get("text", "")))
        elif bt == "markdown":
            for sub in _parse_markdown_blocks(b.get("text", "")):
                render(sub)
        elif bt == "table":
            cols = b.get("columns", [])
            rows = b.get("rows", [])
            if cols:
                tbl_n["n"] += 1
                a(r"\vspace{0.5em}\noindent\textbf{Table %d. %s}\par\vspace{0.3em}"
                  % (tbl_n["n"], _tex_escape(b.get("title", ""))))
                spec = "l" * len(cols)
                a(r"\begin{longtable}{%s}" % spec)
                a(r"\toprule")
                a(" & ".join(_tex_escape(c) for c in cols) + r" \\")
                a(r"\midrule")
                for row in rows:
                    cells = ["" if v is None else _tex_escape(v) for v in row[: len(cols)]]
                    a(" & ".join(cells) + r" \\")
                a(r"\bottomrule")
                a(r"\end{longtable}")
        elif bt == "figure":
            p = b.get("path")
            if p and os.path.exists(p):
                fig_n["n"] += 1
                cap = (b.get("caption", "") or "").replace("Figure: ", "")
                a(r"\begin{figure}[h]\centering")
                a(r"\includegraphics[width=0.8\textwidth]{%s}" % p)
                a(r"\caption{%s}" % _tex_escape(cap))
                a(r"\end{figure}")
        elif bt == "references":
            a(r"\section*{References}")
            a(r"\begingroup\setlength{\parindent}{-0.5in}\setlength{\leftskip}{0.5in}")
            for item in b.get("items", []):
                a(_tex_escape(item) + r"\par")
            a(r"\endgroup")

    for b in blocks:
        render(b)

    a(r"\end{document}")
    _ensure_dir(path)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path


# ----------------------------------------------------------------------------
# HTML (for public shareable read-only report links)
# ----------------------------------------------------------------------------
import base64 as _base64
import html as _html


def _h(text: str) -> str:
    return _html.escape("" if text is None else str(text))


def _html_inline(text: str) -> str:
    out = []
    for seg, bold, ital in _inline_segments(text or ""):
        s = _h(seg)
        if bold:
            s = f"<strong>{s}</strong>"
        if ital:
            s = f"<em>{s}</em>"
        out.append(s)
    return "".join(out)


def _img_data_uri(path: str) -> str | None:
    try:
        with open(path, "rb") as fh:
            data = _base64.b64encode(fh.read()).decode("ascii")
        ext = os.path.splitext(path)[1].lstrip(".").lower() or "png"
        mime = "image/svg+xml" if ext == "svg" else f"image/{ext}"
        return f"data:{mime};base64,{data}"
    except OSError:
        return None


def render_html_body(meta: dict, blocks: list[dict]) -> str:
    """Render report blocks to a self-contained HTML body string.

    Figures are inlined as base64 data URIs so the output needs no external
    files and can be served safely on a public, read-only page.
    """
    parts: list[str] = []
    a = parts.append
    a(f'<header class="rpt-head"><h1>{_h(meta.get("title", "Research Report"))}</h1>')
    if meta.get("subtitle"):
        a(f'<p class="rpt-sub">{_h(meta["subtitle"])}</p>')
    byline = " · ".join(_h(x) for x in (meta.get("author"), meta.get("field"), meta.get("date")) if x)
    if byline:
        a(f'<p class="rpt-by">{byline}</p>')
    a("</header>")

    tbl_n = {"n": 0}
    fig_n = {"n": 0}

    def render(b: dict) -> None:
        bt = b.get("type")
        if bt == "heading":
            lvl = min(3, b.get("level", 2)) + 1  # h2/h3/h4 under the title
            a(f"<h{lvl}>{_h(b.get('text',''))}</h{lvl}>")
        elif bt == "paragraph":
            a(f"<p>{_html_inline(b.get('text',''))}</p>")
        elif bt == "bullet":
            a(f"<ul><li>{_html_inline(b.get('text',''))}</li></ul>")
        elif bt == "number":
            a(f"<ol><li>{_html_inline(b.get('text',''))}</li></ol>")
        elif bt == "markdown":
            for sub in _parse_markdown_blocks(b.get("text", "")):
                render(sub)
        elif bt == "table":
            cols = b.get("columns", [])
            rows = b.get("rows", [])
            if cols:
                tbl_n["n"] += 1
                a(f'<p class="rpt-cap"><strong>Table {tbl_n["n"]}. {_h(b.get("title",""))}</strong></p>')
                a("<table><thead><tr>")
                for c in cols:
                    a(f"<th>{_h(c)}</th>")
                a("</tr></thead><tbody>")
                for row in rows:
                    a("<tr>")
                    for v in row[: len(cols)]:
                        a(f"<td>{_h('' if v is None else v)}</td>")
                    a("</tr>")
                a("</tbody></table>")
        elif bt == "figure":
            uri = _img_data_uri(b.get("path", "")) if b.get("path") else None
            if uri:
                fig_n["n"] += 1
                cap = (b.get("caption", "") or "").replace("Figure: ", "")
                a(f'<figure><img alt="figure" src="{uri}" /><figcaption>Figure {fig_n["n"]}. {_h(cap)}</figcaption></figure>')
        elif bt == "references":
            a("<h2>References</h2><div class='rpt-refs'>")
            for item in b.get("items", []):
                a(f"<p>{_h(item)}</p>")
            a("</div>")

    for b in blocks:
        render(b)
    return "".join(parts)


_SHARED_CSS = """
*{box-sizing:border-box}body{margin:0;background:#f5f7f8;color:#0f1e1c;font-family:ui-sans-serif,system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;line-height:1.6}
.rpt-wrap{max-width:820px;margin:0 auto;padding:40px 20px}
.rpt-card{background:#fff;border:1px solid #e3e8e7;border-radius:16px;padding:48px 56px;box-shadow:0 1px 3px rgba(0,0,0,.04)}
.rpt-head h1{font-size:28px;line-height:1.25;margin:0 0 6px}.rpt-sub{font-style:italic;color:#5b6b68;margin:0 0 4px}.rpt-by{color:#5b6b68;font-size:14px;margin:0 0 8px}
h2{font-size:20px;margin:28px 0 8px;border-bottom:1px solid #e3e8e7;padding-bottom:4px}h3{font-size:17px;margin:20px 0 6px}h4{font-size:15px;margin:16px 0 4px}
p{margin:0 0 12px}ul,ol{margin:0 0 12px 22px}
table{border-collapse:collapse;width:100%;margin:8px 0 20px;font-size:14px}th,td{border:1px solid #e3e8e7;padding:7px 10px;text-align:left}th{background:#f0f4f3}
.rpt-cap{margin:18px 0 4px;font-size:14px}figure{margin:18px 0;text-align:center}figure img{max-width:100%;border:1px solid #e3e8e7;border-radius:8px}figcaption{font-size:13px;color:#5b6b68;font-style:italic;margin-top:6px}
.rpt-refs p{padding-left:2em;text-indent:-2em}
.rpt-foot{max-width:820px;margin:16px auto 0;text-align:center;color:#8a9794;font-size:12px}
"""


def shared_report_page(title: str, body_html: str) -> str:
    """Wrap a rendered body in a full, standalone HTML document."""
    return (
        "<!doctype html><html lang='en'><head><meta charset='utf-8'>"
        "<meta name='viewport' content='width=device-width,initial-scale=1'>"
        "<meta name='robots' content='noindex'>"
        f"<title>{_h(title)}</title><style>{_SHARED_CSS}</style></head>"
        f"<body><div class='rpt-wrap'><article class='rpt-card'>{body_html}</article>"
        "<p class='rpt-foot'>Shared via ResearchAI · read-only</p></div></body></html>"
    )
